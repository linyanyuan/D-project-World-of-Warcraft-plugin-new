"""One-shot generator for the beginner Chinese DOCX manual."""

from __future__ import annotations

from pathlib import Path
from typing import Any

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

OUT = Path(__file__).resolve().parent / "CleanClient-小白完整使用手册.docx"
RELEASE_DOCS = (
    Path(__file__).resolve().parents[1] / "release" / "CleanClient" / "docs"
)


def set_run(
    run: Any, size: float = 11, bold: bool = False, color: RGBColor | None = None
) -> None:
    run.bold = bold
    run.font.size = Pt(size)
    run.font.name = "微软雅黑"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
    if color is not None:
        run.font.color.rgb = color


def add_title(doc: Any, text: str) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(text)
    set_run(r, 22, True)
    p.paragraph_format.space_after = Pt(6)


def add_center(
    doc: Any, text: str, size: float = 11, color: RGBColor | None = None
) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(text)
    set_run(r, size, False, color)
    p.paragraph_format.space_after = Pt(4)


def h(doc, text: str, level: int = 1) -> None:
    p = doc.add_heading(text, level=level)
    for run in p.runs:
        set_run(run, 16 if level == 1 else 13, True)


def para(doc, text: str, bold: bool = False, size: float = 11) -> None:
    p = doc.add_paragraph()
    run = p.add_run(text)
    set_run(run, size, bold)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE


def bullets(doc, items: list[str]) -> None:
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        run = p.add_run(item)
        set_run(run, 11)
        p.paragraph_format.space_after = Pt(2)


def numbered(doc, items: list[str]) -> None:
    for item in items:
        p = doc.add_paragraph(style="List Number")
        run = p.add_run(item)
        set_run(run, 11)
        p.paragraph_format.space_after = Pt(2)


def table(doc, headers: list[str], rows: list[list[str]]) -> None:
    t = doc.add_table(rows=1 + len(rows), cols=len(headers))
    t.style = "Table Grid"
    for i, header in enumerate(headers):
        cell = t.rows[0].cells[i]
        cell.text = ""
        run = cell.paragraphs[0].add_run(header)
        set_run(run, 10.5, True)
    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            cell = t.rows[ri + 1].cells[ci]
            cell.text = ""
            run = cell.paragraphs[0].add_run(val)
            set_run(run, 10.5)
    doc.add_paragraph()


def path_line(doc, text: str) -> None:
    p = doc.add_paragraph()
    run = p.add_run(text)
    set_run(run, 10.5)
    run.font.name = "Consolas"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.left_indent = Cm(0.5)


def build() -> Path:
    doc = Document()
    sec = doc.sections[0]
    sec.top_margin = Cm(2.2)
    sec.bottom_margin = Cm(2.2)
    sec.left_margin = Cm(2.4)
    sec.right_margin = Cm(2.4)

    add_title(doc, "CleanClient 小白完整使用手册")
    add_center(
        doc,
        "魔兽世界自动循环客户端（研究 / 开发版）",
        12,
        RGBColor(0x55, 0x55, 0x55),
    )
    add_center(
        doc,
        "面向完全没有技术背景的使用者 · 按步骤照做即可",
        10.5,
        RGBColor(0x77, 0x77, 0x77),
    )
    add_center(
        doc,
        r"文档位置：docs\CleanClient-小白完整使用手册.docx",
        10,
        RGBColor(0x88, 0x88, 0x88),
    )

    para(doc, "先看这三句（非常重要）", bold=True)
    bullets(
        doc,
        [
            "这个软件现在是「研究 / 开发版」，默认只记日志，不会替你按键打怪。",
            r"发版包里有两样东西：电脑软件 CleanClient.exe，以及魔兽插件 addon\AutoPlayer。"
            "插件不会自动装进游戏，必须你自己复制。",
            "如果日志里一直显示「空闲」，在默认「模拟识别」模式下是正常现象，不代表软件坏了。",
        ],
    )

    h(doc, "一、你需要准备什么")
    numbered(
        doc,
        [
            "一台已经能正常登录《魔兽世界》正式服 / 时光服的电脑。",
            "一份完整的 CleanClient 发版文件夹（不要只有单独一个 exe）。",
            "知道自己魔兽安装在哪个盘（后面要复制插件）。",
        ],
    )
    para(doc, "发版文件夹通常在：")
    path_line(doc, r"D:\project\World of Warcraft plugin new\release\CleanClient")
    para(doc, "打开后，正常应看到这些内容：")
    table(
        doc,
        ["名称", "是什么", "你要做什么"],
        [
            ["CleanClient.exe", "电脑上的控制软件", "双击打开"],
            ["_internal", "软件运行依赖", "不要删、不要单独挪走"],
            [r"addon\AutoPlayer", "魔兽世界插件", "复制到魔兽 AddOns 目录"],
            ["docs", "说明文档", "不会用时打开查看"],
            ["请先读-安装说明.txt", "最短安装提醒", "建议先看一遍"],
        ],
    )

    h(doc, "二、第一步：把插件装进魔兽世界")
    para(
        doc,
        "软件本身不会自动把插件装进游戏。你必须先完成这一步，以后要做真实识别才有基础。",
    )
    para(doc, "1）找到发版包里的插件文件夹", bold=True)
    path_line(doc, r"...\CleanClient\addon\AutoPlayer")
    para(doc, "确认这个文件夹里能看到 AutoPlayer.toc、AutoPlayer.lua 等文件。")

    para(doc, "2）找到魔兽正式服插件目录", bold=True)
    para(doc, "正式服 / 时光服一般是：")
    path_line(doc, r"...\World of Warcraft\_retail_\Interface\AddOns")
    para(doc, "如果你不知道魔兽装在哪：")
    bullets(
        doc,
        [
            "打开战网 → 魔兽世界 → 设置（齿轮）→ 查看安装文件夹。",
            "然后进入：_retail_ → Interface → AddOns。",
            "如果没有 Interface 或 AddOns 文件夹，可以自己新建同名文件夹。",
        ],
    )

    para(doc, "3）复制插件", bold=True)
    numbered(
        doc,
        [
            "先完全退出魔兽世界（不要只退到选人界面）。",
            "把 AutoPlayer 整个文件夹复制到 AddOns 目录下。",
            "复制完成后路径应类似下面这样。",
        ],
    )
    path_line(
        doc,
        r"...\World of Warcraft\_retail_\Interface\AddOns\AutoPlayer\AutoPlayer.toc",
    )
    para(doc, "常见错误：多套了一层文件夹。下面这种是错的：")
    path_line(doc, r"...\AddOns\AutoPlayer\AutoPlayer\AutoPlayer.toc")

    para(doc, "4）在游戏里启用插件", bold=True)
    numbered(
        doc,
        [
            "重新打开魔兽世界。",
            "在选角色界面点击「插件」。",
            "勾选 AutoPlayer。",
            "进入游戏。若插件正常，通常会出现协议相关的色块 / 像素条（具体外观以插件为准）。",
        ],
    )

    h(doc, "三、第二步：打开 CleanClient 软件")
    para(doc, "方式一：双击 exe（推荐小白使用）", bold=True)
    path_line(
        doc,
        r"D:\project\World of Warcraft plugin new\release\CleanClient\CleanClient.exe",
    )
    bullets(
        doc,
        [
            "必须保持整个 CleanClient 文件夹完整。",
            "不要只把 CleanClient.exe 单独拷到桌面再打开，容易打不开或报错。",
            "如果 Windows 提示未知发布者 / 杀毒拦截，可先允许本次运行（本软件是本地打包的 Python 程序）。",
        ],
    )

    para(doc, "方式二：用源码启动（给会敲命令的人）", bold=True)
    path_line(doc, r'cd "D:\project\World of Warcraft plugin new"')
    path_line(doc, r"pip install -r clean_client\requirements.txt")
    path_line(doc, "python -m clean_client.app")

    h(doc, "四、软件界面怎么看")
    para(doc, "左侧一般有四个页面：控制台、循环、识别、系统设置。")

    h(doc, "1. 控制台（最常用）", 2)
    table(
        doc,
        ["你看到的文字", "作用", "小白建议"],
        [
            ["启动", "开始后台循环", "确认「只记日志」勾选后再点"],
            ["停止", "结束循环", "不用时点停止"],
            ["只记日志，不按键", "只写日志，不向游戏发按键", "默认勾着，务必先保持勾选"],
            ["优先高亮技能", "识别时优先考虑高亮技能", "可先保持默认"],
            ["截屏方式", "怎么抓游戏画面", "第一次建议选「空（测试）」"],
            ["周期(ms)", "循环间隔毫秒", "默认约 30，先别改"],
            ["运行状态 / 待机 / 运行中", "当前是否在跑", "点启动后应变为运行中"],
            ["窗口句柄", "有没有找到魔兽窗口", "找不到也可能先空着，后面再查"],
            ["运行日志", "软件输出信息", "重点看这里有没有「空闲 / 动作 / 错误」"],
        ],
    )

    h(doc, "2. 循环", 2)
    para(
        doc,
        "这里显示当前加载的技能优先级列表（例如邪恶死亡骑士）。当前版本主要是查看，一般不用改。",
    )

    h(doc, "3. 识别", 2)
    table(
        doc,
        ["选项", "含义", "现在怎么选"],
        [
            ["模拟识别（调试）", "不读屏幕，假装没有技能", "当前默认，试用时用这个"],
            ["像素协议（真实色块）", "读取 AutoPlayer 画出的色块", "以后接好真实识别再用"],
            ["区域目录", "Skill/Target/Player/Buff 区域文件所在文件夹", "标定保存后可在这里指定"],
            ["打开标定器", "弹出框选区域的工具", "想练习框选时点它"],
        ],
    )

    h(doc, "4. 系统设置", 2)
    bullets(
        doc,
        [
            "窗口关键字：默认包含 World of Warcraft、魔兽世界，用于查找游戏窗口。",
            "冷却就绪窗口、增益匹配阈值：高级参数，小白先保持默认。",
            "改完后如有保存按钮 / 提示，按界面提示保存到本地配置。",
        ],
    )

    h(doc, "五、第一次安全试用（强烈建议按这个做）")
    para(doc, "目标：确认软件能打开、能启动、能出日志。不追求自动打怪。")
    numbered(
        doc,
        [
            "确认魔兽可以正常进入（插件已勾选更好，但第一次即使没装插件也能先测软件本身）。",
            "打开 CleanClient.exe。",
            "进入「控制台」。",
            "确认已勾选「只记日志，不按键」。",
            "截屏方式先选「空（测试）」。",
            "点「启动」。",
            "看运行日志：多半会出现「空闲」。这表示循环在跑，但当前是模拟识别，没有读到真实技能。",
            "点「停止」结束。",
        ],
    )
    para(doc, "只要能启动、能停止、日志有输出，就说明软件本体工作正常。", bold=True)

    h(doc, "六、截屏方式怎么选")
    table(
        doc,
        ["界面选项", "简单理解", "什么时候用"],
        [
            ["空（测试）", "不真正截屏", "第一次试用、排查界面"],
            ["方式一", "按窗口内容抓图（PrintWindow）", "窗口模式可优先试"],
            ["方式二", "抓屏幕上实际看到的画面（MSS）", "一般最稳妥，常作首选"],
            ["方式三", "更快的桌面复制（DXGI/dxcam）", "方式二不行或想更快时再试"],
        ],
    )
    bullets(
        doc,
        [
            "某种方式黑屏、报错、抓不到图：换另一种即可。",
            "全屏独占、HDR、多显示器有时会导致某种方式失败，这属于环境差异，不是你操作错了。",
            "做区域标定时，尽量让截屏方式和以后正式使用时一致。",
        ],
    )

    h(doc, "七、区域标定（给以后真实识别做准备）")
    para(doc, "标定的意思：在游戏画面上框出软件要盯着看的几个区域，例如技能条色块区。")
    para(doc, "打开方式：")
    bullets(
        doc,
        [
            "在软件「识别」页点击「打开标定器」；或",
            "命令行运行：python -m clean_client.tools.calibrate_regions",
        ],
    )
    para(doc, "建议练习流程：")
    numbered(
        doc,
        [
            "截屏先选「空（测试）」，点「抓取画面」，先熟悉界面。",
            "选择区域类型：技能 / 目标 / 玩家 / 增益。",
            "在图上按住鼠标拖拽框选；也可以改左侧数字坐标。",
            "点「保存区域」，选一个文件夹保存。",
            "保存后通常会生成 skill_region.txt、target_region.txt、player_region.txt、buff_region.txt。",
        ],
    )
    para(
        doc,
        "实机标定时：先打开魔兽并启用 AutoPlayer，再用方式一/二/三抓真实画面，重点框准技能（Skill）色块区域。",
    )
    para(
        doc,
        "说明：标定文件是为后续真实识别准备的。若当前版本主界面还没有完全自动读取你的标定目录，属于功能尚未接完，不是你保存失败。",
        bold=True,
    )

    h(doc, "八、现在做得到 / 做不到什么")
    table(
        doc,
        ["事项", "当前状态"],
        [
            ["打开中文界面", "可以"],
            ["启动 / 停止循环", "可以"],
            ["只记日志，不按键", "可以（默认开启）"],
            ["查看循环技能列表", "可以"],
            ["练习区域标定", "可以"],
            ["自动打怪 / 自动按键输出", "默认不会；真实识别也尚未完全接好"],
            ["免登录使用", "可以（暂无卡密）"],
        ],
    )

    h(doc, "九、发给朋友时怎么发")
    numbered(
        doc,
        [
            "发送整个 CleanClient 文件夹（可打成 zip）。",
            "让朋友先看「请先读-安装说明.txt」和本手册。",
            r"朋友必须先把 addon\AutoPlayer 复制到自己的魔兽 AddOns。",
            "再运行 CleanClient.exe。",
        ],
    )
    para(
        doc,
        "如果仓库是 GitHub 私有库：朋友需要你把其加成协作者后，才能下载 Releases。",
    )

    h(doc, "十、常见问题")
    faq = [
        (
            "Q1：为什么打包目录里一开始找不到 AutoPlayer？",
            r"A：旧版打包可能只打了 exe。现在新打包会把 addon\AutoPlayer、说明文档一起放进 release\CleanClient。"
            r"若你手里的包没有，请重新运行 scripts\pack-CleanClient.bat。",
        ),
        (
            "Q2：双击打包脚本没反应？",
            r"A：请改双击 scripts\pack-CleanClient.bat。正常会弹出黑窗口运行约 1.5～2 分钟。若秒退，把窗口里的报错发出来。",
        ),
        (
            "Q3：日志一直是「空闲」？",
            "A：默认「模拟识别」不读屏，所以一直空闲是正常的。",
        ),
        (
            "Q4：会不会给我乱按键？",
            "A：勾选「只记日志，不按键」时不会。请一直保持勾选，直到你明确知道自己在做什么。",
        ),
        (
            "Q5：插件复制了但游戏里没有？",
            r"A：检查是否复制到 _retail_\Interface\AddOns；是否多套一层文件夹；是否在选人界面启用；是否完全重启过游戏。",
        ),
        (
            "Q6：杀毒软件报毒？",
            "A：本地 PyInstaller 打包程序常见误报。请先确认文件来源是你自己的项目打包结果，再加入信任区。",
        ),
        (
            "Q7：我想要真正自动循环，下一步要做什么？",
            "A：1）装好并启用 AutoPlayer；2）用真实截屏标定 Skill 等区域；3）把识别模式改到像素协议；"
            "4）确认日志里能看到真实动作后，再考虑取消「只记日志」。当前版本尚未把这条链路完全打通。",
        ),
    ]
    for q, a in faq:
        para(doc, q, bold=True)
        para(doc, a)

    h(doc, "十一、推荐操作清单（打勾用）")
    bullets(
        doc,
        [
            "已拿到完整 CleanClient 文件夹",
            r"已把 addon\AutoPlayer 复制到 _retail_\Interface\AddOns\AutoPlayer",
            "游戏选人界面已勾选 AutoPlayer",
            "已双击 CleanClient.exe 打开软件",
            "控制台已勾选「只记日志，不按键」",
            "截屏方式先选「空（测试）」并成功启动 / 停止",
            "日志能看到「空闲」或其他输出",
            "（可选）已打开标定器练习框选并保存",
        ],
    )

    h(doc, "十二、相关文件位置")
    table(
        doc,
        ["内容", "路径"],
        [
            ["发版目录", r"release\CleanClient"],
            ["本手册", r"docs\CleanClient-小白完整使用手册.docx"],
            ["简版安装说明", r"release\CleanClient\请先读-安装说明.txt"],
            ["Markdown 手册", r"docs\使用手册-clean_client.md"],
            ["重新打包", r"scripts\pack-CleanClient.bat"],
            ["插件源目录", r"addon\AutoPlayer"],
        ],
    )

    para(
        doc,
        "—— 文档结束 —— 若某一步和你屏幕上不一致，把截图发出来即可对照修改。",
        bold=True,
    )

    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(OUT))
    if RELEASE_DOCS.exists():
        RELEASE_DOCS.mkdir(parents=True, exist_ok=True)
        doc.save(str(RELEASE_DOCS / OUT.name))
    return OUT


if __name__ == "__main__":
    path = build()
    print(path)
    print("size=", path.stat().st_size)
