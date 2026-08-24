"""Generate Chinese CleanClient user manual as .docx."""

from __future__ import annotations

from pathlib import Path
from typing import Any

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


def _set_run_font(
    run, *, size: float = 12, bold: bool = False, color: RGBColor | None = None
) -> None:
    run.bold = bold
    run.font.size = Pt(size)
    run.font.name = "微软雅黑"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
    if color is not None:
        run.font.color.rgb = color


def _add_heading(doc: Any, text: str, level: int = 1) -> None:
    p = doc.add_heading(text, level=level)
    for run in p.runs:
        _set_run_font(run, size=16 if level == 1 else 14, bold=True)


def _add_para(doc: Any, text: str, *, bold: bool = False, size: float = 11) -> None:
    p = doc.add_paragraph()
    run = p.add_run(text)
    _set_run_font(run, size=size, bold=bold)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = 1.35


def _add_bullets(doc: Any, items: list[str]) -> None:
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        run = p.add_run(item)
        _set_run_font(run, size=11)
        p.paragraph_format.space_after = Pt(3)


def _add_table(doc: Any, headers: list[str], rows: list[list[str]]) -> None:
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = ""
        run = cell.paragraphs[0].add_run(h)
        _set_run_font(run, size=10.5, bold=True)
    for r_idx, row in enumerate(rows):
        for c_idx, value in enumerate(row):
            cell = table.rows[r_idx + 1].cells[c_idx]
            cell.text = ""
            run = cell.paragraphs[0].add_run(value)
            _set_run_font(run, size=10.5)
    doc.add_paragraph()


def build_doc(out_path: Path) -> Path:
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Cm(2.2)
    section.bottom_margin = Cm(2.2)
    section.left_margin = Cm(2.4)
    section.right_margin = Cm(2.4)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("CleanClient 中文使用说明")
    _set_run_font(run, size=22, bold=True)

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = sub.add_run("自动循环客户端（研究 / 开发版）")
    _set_run_font(run, size=12, color=RGBColor(0x55, 0x55, 0x55))

    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = meta.add_run("适用目录：D:\\project\\World of Warcraft plugin new")
    _set_run_font(run, size=10.5, color=RGBColor(0x66, 0x66, 0x66))

    _add_para(
        doc,
        "当前版本特点：免登录；默认「只记日志，不按键」；默认「模拟识别（不读屏）」。"
        "因此现在点「启动」通常只会看到「空闲」，不会自动打怪——这是正常现象。",
    )

    _add_heading(doc, "一、如何打开软件", 1)
    _add_para(doc, "方式 A：双击已打包程序（推荐）", bold=True)
    _add_para(
        doc,
        r"D:\project\World of Warcraft plugin new\release\CleanClient\CleanClient.exe",
    )
    _add_bullets(
        doc,
        [
            "请保持整个 CleanClient 文件夹完整，不要只拷贝单个 exe。",
            "重新打包：双击 scripts\\一键打包-CleanClient.bat（精简打包约 1.5～2 分钟）。",
        ],
    )
    _add_para(doc, "方式 B：源码启动（开发）", bold=True)
    _add_para(
        doc,
        'cd "D:\\project\\World of Warcraft plugin new"\n'
        "pip install -r clean_client/requirements.txt\n"
        "python -m clean_client.app",
    )

    _add_heading(doc, "二、界面导览", 1)
    _add_table(
        doc,
        ["页面", "作用"],
        [
            ["控制台", "启动 / 停止、截屏方式、周期、实时日志"],
            ["循环", "查看当前技能优先级列表（只读）"],
            ["识别", "选择模拟识别或像素协议，打开区域标定器"],
            ["系统设置", "窗口关键字、阈值，保存到本地 default.json"],
        ],
    )

    _add_heading(doc, "三、截屏方式说明（方式一 / 二 / 三）", 1)
    _add_para(
        doc,
        "软件需要周期性抓取魔兽窗口画面，才能读取插件画出的色块信息。"
        "「截屏方式」决定用哪一种 Windows 抓图技术：",
    )
    _add_table(
        doc,
        ["界面选项", "技术名", "含义与适用场景"],
        [
            [
                "空（测试·不截屏）",
                "null",
                "不真正截屏，只用于界面联调。日志多为「空闲」。新手先用这个。",
            ],
            [
                "方式一 · 窗口打印（PrintWindow）",
                "PrintWindow",
                "通过 Windows 窗口打印接口抓指定窗口内容。窗口模式较常用；被遮挡时有时仍能抓到。",
            ],
            [
                "方式二 · 屏幕截取（MSS）",
                "MSS",
                "按屏幕区域截图，抓的是屏幕上实际显示的画面。简单稳定，适合大多数情况。",
            ],
            [
                "方式三 · 高速复制（DXGI）",
                "dxcam / DXGI",
                "用 DirectX 桌面复制，通常更快、延迟更低。追求流畅时可优先试这个。",
            ],
        ],
    )
    _add_bullets(
        doc,
        [
            "如果某种方式黑屏 / 报错：换另外一种即可。",
            "全屏独占、多显示器、HDR 可能导致个别方式失败，这是系统限制，不是软件坏了。",
            "标定区域时，截屏方式尽量与正式运行时保持一致。",
        ],
    )

    _add_heading(doc, "四、最小试用流程（不碰游戏）", 1)
    _add_bullets(
        doc,
        [
            "打开 CleanClient.exe。",
            "在「控制台」确认已勾选「只记日志，不按键」。",
            "截屏方式先选「空（测试·不截屏）」。",
            "点「▶ 启动」，状态变为「● 运行中」，日志出现「空闲」。",
            "点「■ 停止」，回到「● 待机」。",
        ],
    )

    _add_heading(doc, "五、为真自动做准备", 1)
    _add_para(doc, "5.1 安装游戏内插件", bold=True)
    _add_bullets(
        doc,
        [
            r"把项目中的 addon\AutoPlayer 复制到魔兽插件目录，例如：",
            r"World of Warcraft\_retail_\Interface\AddOns\AutoPlayer",
            "进入游戏前勾选插件；分辨率 / UI 缩放尽量固定。",
        ],
    )
    _add_para(doc, "5.2 区域标定", bold=True)
    _add_bullets(
        doc,
        [
            "打开「识别」页 →「打开区域标定器」。",
            "选择截屏方式 →「抓取画面」。",
            "分别框选：技能 / 目标 / 玩家 / 增益（技能最重要）。",
            "「保存区域…」到例如 regions\\ 目录，生成 4 个 *_region.txt。",
        ],
    )
    _add_para(doc, "5.3 识别页设置", bold=True)
    _add_bullets(
        doc,
        [
            "模式：调试用「模拟识别」；真读色块用「像素协议」。",
            "区域目录：浏览并选中你保存的 regions 文件夹。",
        ],
    )
    _add_para(doc, "5.4 系统设置", bold=True)
    _add_bullets(
        doc,
        [
            "窗口关键字建议保留：World of Warcraft, 魔兽世界",
            "保存设置后，回到控制台启动，确认「窗口句柄」显示为数字而不是 —。",
        ],
    )
    _add_para(doc, "5.5 何时取消「只记日志，不按键」", bold=True)
    _add_bullets(
        doc,
        [
            "插件已加载，技能区域已标定。",
            "识别改为「像素协议」。",
            "日志开始出现「动作 法术ID=…」且看起来合理。",
            "再取消勾选。取消后才会向游戏发送按键，请谨慎。",
        ],
    )

    _add_heading(doc, "六、常见问题", 1)
    _add_table(
        doc,
        ["问题", "说明"],
        [
            [
                "为什么一直「空闲」？",
                "模拟识别不读屏；或像素协议下区域未标定 / 目录未选。",
            ],
            ["会不会帮我按键？", "勾选「只记日志，不按键」时不会。"],
            ["要不要登录卡密？", "当前版本不需要。"],
            ["标定了为什么没变化？", "请在「识别」页手动选择 regions 目录。"],
            [
                "打包为什么以前很慢？",
                "旧脚本误打入 torch 等大火腿；现已精简到约 300MB、约 2 分钟。",
            ],
        ],
    )

    _add_heading(doc, "七、相关文件", 1)
    _add_bullets(
        doc,
        [
            r"程序：release\CleanClient\CleanClient.exe",
            r"图标资源：clean_client\assets\cleanclient.ico",
            r"简版说明：clean_client\README.md",
            r"本手册源脚本：docs\generate_user_manual_docx.py",
            r"卡密延期计划：docs\auth-cardkey-plan.md",
        ],
    )

    _add_para(
        doc,
        "声明：全自动操作可能违反游戏服务条款。本软件仅供私人研究与自有账号实验，请自行承担风险。",
        size=10,
    )

    out_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(out_path))
    return out_path


if __name__ == "__main__":
    target = Path(__file__).resolve().parent / "CleanClient-中文使用说明.docx"
    path = build_doc(target)
    print(path)
